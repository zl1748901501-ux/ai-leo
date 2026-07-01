from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\张龙\Desktop\AI产品经理简历")
OUT_PATH = OUT_DIR / "张龙_AI产品经理_简历.docx"


def set_font(run, size=10.5, bold=False, color="1E293B"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_bottom_border(paragraph, color="CBD5E1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 12, True, "1D4ED8")
    add_bottom_border(p, "DBEAFE")


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 9.4, True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, 9.4)
    else:
        set_font(p.add_run(text), 9.4)


def line(doc, left, right="", size=10, bold=False, color="0F172A"):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.allow_autofit = True
    cells = table.rows[0].cells
    for cell in cells:
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tc_pr.append(border)
    r = cells[0].paragraphs[0].add_run(left)
    set_font(r, size, bold, color)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = cells[1].paragraphs[0].add_run(right)
    set_font(r2, size, False, "475569")


def build_resume():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("张龙")
    set_font(run, 22, True, "0F172A")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run("AI 产品经理 / AI 产品实习生 / 2026 届数字媒体艺术应届生"), 10.5, True, "1D4ED8")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        contact.add_run("电话：15937126737  |  微信：zllll_0216  |  邮箱：1748901501@qq.com  |  现居：河南郑州"),
        9.2,
        False,
        "475569",
    )

    heading(doc, "求职优势")
    bullet(
        doc,
        "数字媒体艺术背景，具备视觉表达、作品集包装、用户体验和内容呈现能力，正在转向 AI 产品经理方向。",
    )
    bullet(
        doc,
        "独立推进 Second AI 从产品定位、PRD、原型设计到可运行 MVP 的完整闭环，理解 AI SaaS 的资料上传、AI 解析、权限控制、访客问答和证据展示流程。",
    )
    bullet(
        doc,
        "熟悉用 AI 编程工具协作实现产品原型，能把需求拆成页面、状态、数据表、权限规则和验收标准，并持续根据体验问题迭代。",
    )

    heading(doc, "项目经历")
    line(doc, "Second AI - 隐私优先的个人 AI Workspace / 可被提问的 AI 作品主页", "2026.06 - 2026.07", 10.2, True)
    bullet(
        doc,
        "项目定位：面向个人用户的 AI SaaS 原型，用户上传简历、PDF、Word、图片、视频和作品资料后，AI 进行理解、摘要、标签和问答；访客通过受控邀请进入 AI Profile，与 AI 对话了解资料主人。",
        "项目定位：",
    )
    bullet(
        doc,
        "核心问题：传统网盘只存储不理解，普通聊天机器人缺少长期资料管理和受控展示；Second AI 尝试把资料管理、AI 理解、作品集展示和访客问答组合成一个个人 AI 工作空间。",
        "核心问题：",
    )
    bullet(
        doc,
        "主要工作：完成产品定位、MVP 范围、页面信息架构、Owner/Visitor 双端流程、权限规则、邀请 Token 校验、资料状态和 AI 问答边界设计。",
        "主要工作：",
    )
    bullet(
        doc,
        "功能闭环：实现登录注册、Dashboard、上传中心、资产库、资产详情、AI 主页生成、访问权限、访客 AI 主页、无权限页等核心页面。",
        "功能闭环：",
    )
    bullet(
        doc,
        "AI 能力：接入 DeepSeek 进行资料解析和访客问答；设计 private / answer_only / show 三档权限，确保 AI 只基于授权资料回答，并展示来源引用和相关资料卡片。",
        "AI 能力：",
    )
    bullet(
        doc,
        "产品思考：围绕隐私、访问控制、资料可信来源、访客体验和 AI 不编造原则持续迭代，形成“可被提问的 AI 作品主页”差异化表达。",
        "产品思考：",
    )
    bullet(
        doc,
        "技术协作：使用 Next.js、TypeScript、Tailwind CSS、Supabase Auth/Database/Storage、DeepSeek API，借助 AI 编程工具完成可演示版本并进行多轮验收。",
        "技术协作：",
    )

    line(doc, "数字媒体作品与三维场景实践", "2021 - 2026", 10.2, True)
    bullet(doc, "围绕场景建模、场景原画、人物设计、剪辑、数字虚拟空间可视化等方向完成课程与作品实践。")
    bullet(doc, "熟悉 3D 建模、PBR 流程、UE 场景搭建和视觉表达，能够将视觉叙事能力迁移到 AI 产品原型、作品展示和交互体验设计中。")

    heading(doc, "教育背景")
    line(doc, "川音成都美术学院 | 数字媒体艺术 | 本科", "2024.09 - 2026.06", 10.1, True)
    bullet(doc, "主修课程：数字项目设计与制作、数字虚拟空间可视化、项目策划、数字产品创新型设计等。")
    line(doc, "四川华新现代职业技术学院 | 数字艺术 | 专科", "2021.09 - 2024.06", 10.1, True)
    bullet(doc, "主修课程：场景建模、场景原画、人物设计、剪辑等。")

    heading(doc, "实践经历")
    bullet(doc, "校内实践：参与工作室 Logo、海报宣发、展板设计；参与创建校内比赛“鹦鹉螺杯”。")
    bullet(doc, "活动执行：参与开学典礼、元旦晚会等活动策划、视觉物料、场景布置和人员分工。")
    bullet(doc, "校外实践：参与四川省青少年语言艺术展示活动比赛项目会场执行；参与颂果传媒拍摄、剪辑、策划和人员分配工作。")

    heading(doc, "技能")
    bullet(doc, "产品能力：PRD 撰写、信息架构、用户流程、权限设计、原型设计、功能验收、AI 产品场景拆解。")
    bullet(doc, "AI/技术理解：DeepSeek API、Supabase Auth/Database/Storage、Next.js 原型协作、AI 解析、访客问答、OCR/视觉识别边界。")
    bullet(doc, "设计与内容：作品集包装、视觉表达、交互原型、PS、AI、PR、AE。")
    bullet(doc, "三维与引擎：3ds Max、ZBrush、TopoGun、Substance 3D、Marmoset Toolbag、Unreal Engine。")

    heading(doc, "荣誉奖项")
    bullet(doc, "2023 香港大学生当代设计银奖、铜奖。")
    bullet(doc, "2022-2023 第五届日本概念艺术设计银奖、铜奖。")
    bullet(doc, "2023 ROCA 平面与空间设计金级、铜奖、优秀奖。")
    bullet(doc, "2022-2023 国家励志奖学金。")

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_resume())
