from __future__ import annotations

import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DESKTOP = Path(r"C:\Users\张龙\Desktop")
OUT_DIR = DESKTOP / "ai设计稿"
DOC_DIR = OUT_DIR / "文档"
IMG_DIR = OUT_DIR / "原型截图"
SRC_IMG_DIR = DESKTOP / "原型ai图"

TODAY = date.today().strftime("%Y.%m.%d")


SCREENSHOTS = [
    ("01-login-登录页.png", "登录页", "用户输入邮箱和密码进入工作台；新账号可自动注册，降低初期使用门槛。"),
    ("02-dashboard-工作台.png", "Workspace 工作台", "Owner 端总览资料状态、资料查找、下载申请、隐私边界和最近资产。"),
    ("03-upload-上传中心.png", "上传中心", "支持拖拽或点击上传，上传后写入资产表，并进入 AI 解析流程。"),
    ("04-assets-prd-资产详情.png", "资产详情", "展示文件预览、AI 解析结果、可回答问题和权限设置。"),
    ("05-profile-generator-AI主页生成.png", "AI 主页生成", "配置 AI Profile 的身份信息、简介、推荐问题和可展示资料。"),
    ("06-share-访问权限.png", "访问权限", "Owner 输入访客邮箱生成邀请链接，支持复制、撤销和有效期设置。"),
    ("07-visitor-isabella-访客AI主页.png", "访客 AI 主页", "访客通过受控邀请进入，与 AI 对话并查看授权资料。"),
    ("08-no-access-无权限页.png", "无权限页", "邀请无效、过期、邮箱不匹配或未登录时展示访问失败状态。"),
]


def ensure_dirs() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    for name, _, _ in SCREENSHOTS:
        src = SRC_IMG_DIR / name
        if src.exists():
            shutil.copy2(src, IMG_DIR / name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1D4ED8"),
        ("Heading 2", 13, "334155"),
        ("Heading 3", 11, "475569"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(71, 85, 105)

    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"版本：V0.1 初期稿    日期：{TODAY}    项目：Second AI")
    r3.font.name = "Microsoft YaHei"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(100, 116, 139)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(29, 78, 216)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def build_prd() -> Path:
    doc = Document()
    set_doc_styles(doc)
    add_title(
        doc,
        "Second AI 初期产品 PRD",
        "隐私优先的个人 AI Workspace 与可被提问的 AI 作品主页",
    )
    add_callout(
        doc,
        "一句话定位",
        "上传个人资料，让 AI 理解、管理、检索、展示，并在受控权限下替资料主人回答访客问题。",
    )

    doc.add_heading("1. 产品背景", level=1)
    doc.add_paragraph(
        "Second AI 的初期目标不是做一个普通聊天机器人，也不是传统网盘，而是把“资料管理、AI 理解、受控分享、访客问答”组合成一个轻量可用的个人 AI Workspace。"
    )
    doc.add_paragraph(
        "用户可以上传简历、PDF、Word、图片、作品说明等资料，系统完成基础解析后，将这些资料变成 AI 可理解的资产。访客进入受控主页后，不再浏览静态作品集，而是通过提问了解资料主人。"
    )

    doc.add_heading("2. 初期目标", level=1)
    add_bullets(
        doc,
        [
            "完成从注册登录、上传资料、AI 解析、权限设置、生成邀请到访客问答的最小闭环。",
            "Owner 端可以管理资料，控制哪些资料私密、仅回答或可展示。",
            "Visitor 端必须通过邀请访问，AI 只能基于授权资料回答。",
            "第一版先追求产品闭环与可演示性，不做复杂 RAG、团队版、支付、多租户和 Agent。",
        ],
    )

    doc.add_heading("3. 目标用户与场景", level=1)
    add_table(
        doc,
        ["角色", "核心需求", "初期场景"],
        [
            ["资料主人 Owner", "管理个人资料并控制展示边界", "上传简历、PRD、作品说明、图片/PDF，设置权限并生成邀请链接。"],
            ["访客 Visitor", "快速了解资料主人能力与作品", "通过受邀链接登录访问，向 AI 提问并查看授权资料。"],
            ["面试官/合作方", "高效判断资料主人是否匹配", "询问经历、项目、能力、作品，并查看相关证据卡片。"],
        ],
        [1.3, 2.1, 3.0],
    )

    doc.add_heading("4. 核心功能范围", level=1)
    add_table(
        doc,
        ["模块", "初期能力", "说明"],
        [
            ["登录注册", "邮箱密码登录，新账号自动注册", "降低试用门槛；后续可补邮箱验证和第三方登录。"],
            ["上传中心", "上传文件到 Supabase Storage", "支持 PDF、Word、Markdown、TXT、图片、视频等基础格式。"],
            ["AI 解析", "提取正文并生成标题、摘要、标签、问题、卡片", "DeepSeek 负责文本解析；图片/扫描 PDF 需单独 OCR/视觉服务。"],
            ["资产库", "搜索、查看、删除资料", "Owner 可以快速定位个人资料。"],
            ["资产权限", "私密 / 仅回答 / 可展示", "控制资料是否进入访客问答、是否展示卡片、是否允许下载。"],
            ["AI 主页生成", "配置身份信息、简介、推荐问题", "形成可被提问的个人 AI Portfolio。"],
            ["邀请访问", "生成 token 邀请链接，可撤销、可设置有效期", "访客邮箱与邀请邮箱一致才可访问。"],
            ["访客问答", "基于授权资料回答并展示来源与资料卡片", "不做 embedding，只做授权资料上下文拼接。"],
        ],
        [1.3, 2.1, 3.0],
    )

    doc.add_heading("5. 权限规则", level=1)
    add_table(
        doc,
        ["权限", "AI 是否可用于回答", "访客是否可打开", "访客是否可下载"],
        [
            ["私密 private", "否", "否", "否"],
            ["仅回答 answer_only", "是", "是，可网页打开查看", "否"],
            ["可展示 show", "是", "是，可网页打开查看", "是"],
        ],
        [1.4, 1.7, 1.8, 1.5],
    )

    doc.add_heading("6. 关键页面", level=1)
    add_bullets(
        doc,
        [
            "/login：登录与自动注册入口。",
            "/dashboard：工作台，展示资料概览、资料查找、下载申请和隐私状态。",
            "/upload：上传中心，支持点击和拖拽上传。",
            "/assets：资产库，展示全部已上传资料，可搜索和删除。",
            "/assets/[id]：资产详情，展示文件预览、AI 解析、权限设置。",
            "/profile-generator：AI 主页基础信息配置。",
            "/share：生成邀请链接、复制链接、撤销邀请、设置有效期。",
            "/visitor/isabella：访客 AI Portfolio，展示问答、来源、相关资料。",
            "/no-access：无权限或邀请失效状态。",
        ],
    )

    doc.add_heading("7. 初期验收标准", level=1)
    add_bullets(
        doc,
        [
            "用户可以注册/登录并保持会话。",
            "用户可以上传文件并在资产库看到记录。",
            "用户可以进入资产详情，看到文件预览和 AI 解析结果。",
            "用户可以设置资料权限，并影响访客端资料展示。",
            "Owner 可以生成邀请链接，访客登录后通过校验进入主页。",
            "访客可以提问，AI 只基于授权资料回答，并展示来源和相关资料。",
            "未授权邮箱、撤销邀请、过期邀请均不能访问访客主页。",
        ],
    )

    doc.add_heading("8. 暂不包含", level=1)
    add_bullets(
        doc,
        [
            "复杂 RAG、embedding、pgvector 和全文 chunk 检索。",
            "团队版、多租户、支付、套餐和商业后台。",
            "下载申请真实审批流。",
            "完整访问日志与风控系统。",
            "Agent 自动整理、多轮任务执行。",
        ],
    )

    path = DOC_DIR / "Second AI 初期PRD.docx"
    doc.save(path)
    return path


def build_design_doc() -> Path:
    doc = Document()
    set_doc_styles(doc)
    add_title(
        doc,
        "Second AI 原型与设计过程说明",
        "从 AI Workspace 到受控访问 AI Portfolio 的初期设计稿",
    )
    add_callout(
        doc,
        "设计目标",
        "Owner 端强调清晰、高效、可控；Visitor 端强调沉浸式、作品集感和 AI Profile 的提问体验。",
    )

    doc.add_heading("1. 设计方向", level=1)
    add_table(
        doc,
        ["端", "视觉方向", "设计理由"],
        [
            ["Owner 管理端", "浅色 AI Workspace：深色侧边栏、浅色内容区、圆角卡片、蓝紫/青蓝强调色", "管理资料时需要高效、清晰、稳定，避免过度沉浸影响操作。"],
            ["Visitor 访客端", "暗色沉浸式 AI Portfolio：蓝紫渐变、对话核心、相关资料侧栏", "访客不是进后台，而是在提问一个可交互的作品主页。"],
        ],
        [1.3, 2.5, 2.6],
    )

    doc.add_heading("2. 核心交互流程", level=1)
    add_bullets(
        doc,
        [
            "Owner 注册/登录进入 Workspace。",
            "Owner 上传资料，系统写入 Storage 和 assets 表。",
            "Owner 对资料进行 AI 解析，生成摘要、标签、可回答问题和资料卡片。",
            "Owner 设置每份资料的权限：私密、仅回答、可展示。",
            "Owner 在访问权限页输入访客邮箱，生成邀请链接并自动复制。",
            "Visitor 打开链接后登录，系统校验 token、邮箱、状态和有效期。",
            "校验通过后进入访客 AI 主页，可以提问并查看授权资料。",
            "AI 回答时展示来源引用和相关资料卡片。",
        ],
    )

    doc.add_heading("3. 页面原型说明", level=1)
    for file_name, title, note in SCREENSHOTS:
        doc.add_heading(title, level=2)
        doc.add_paragraph(note)
        img_path = IMG_DIR / file_name
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    doc.add_heading("4. 关键设计决策", level=1)
    add_bullets(
        doc,
        [
            "访客页不使用普通聊天机器人样式，而是使用 AI Portfolio 表达，让用户感觉是在了解一个人和其作品。",
            "右侧相关资料不只是文件列表，而是回答证据和可打开资料，强化“AI 回答有依据”。",
            "权限设计采用三档，足够覆盖初期隐私需求，也便于访客端展示逻辑落地。",
            "上传与解析分离，先让真实上传可用，再逐步增强 OCR、视觉识别和文档解析能力。",
            "问答先采用轻量上下文，不做复杂 RAG，降低第一版开发复杂度。",
        ],
    )

    doc.add_heading("5. 后续可优化方向", level=1)
    add_bullets(
        doc,
        [
            "增加 OCR 识别结果预览和手动修正，避免图片/扫描 PDF 识别不准。",
            "把下载申请做成真实审批流，并记录访问日志。",
            "引入 embedding 检索和 chunk 级引用，提高资料较多时的问答准确率。",
            "支持更多资料类型，如 GitHub、网页链接、视频字幕和设计工程文件。",
            "完善公开/私密主页配置，让用户可选择完全私密、受邀访问或公开展示。",
        ],
    )

    path = DOC_DIR / "Second AI 原型与设计过程说明.docx"
    doc.save(path)
    return path


def build_readme() -> Path:
    path = OUT_DIR / "README.txt"
    path.write_text(
        "\n".join(
            [
                "Second AI 设计稿资料包",
                f"生成日期：{TODAY}",
                "",
                "包含内容：",
                "1. 文档/Second AI 初期PRD.docx",
                "2. 文档/Second AI 原型与设计过程说明.docx",
                "3. 原型截图/ 8 张页面截图",
                "",
                "说明：本资料包聚焦初期产品闭环，不包含复杂 RAG、团队版、支付、多租户和 Agent 等后期规划。",
            ]
        ),
        encoding="utf-8",
    )
    return path


def package_zip() -> Path:
    zip_path = DESKTOP / "ai设计稿.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in OUT_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(OUT_DIR.parent))
    return zip_path


def main() -> None:
    ensure_dirs()
    prd = build_prd()
    design = build_design_doc()
    readme = build_readme()
    zip_path = package_zip()
    print(prd)
    print(design)
    print(readme)
    print(zip_path)


if __name__ == "__main__":
    main()
